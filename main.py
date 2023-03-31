from fastapi import FastAPI, Request, Header
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse
from fastapi.templating import Jinja2Templates
import base64
from io import BytesIO
from PyPDF2 import PdfReader
import warnings
warnings.simplefilter(action='ignore', category=FutureWarning)
import pandas as pd
import os
import requests
from _md5 import md5
import docx
from azure.storage.blob import ContentSettings

from azure.storage.blob import BlobServiceClient

# Replace this with your connection string
azure_connection_string = "DefaultEndpointsProtocol=https;AccountName=darablobstorage;AccountKey=elaFVu4ns1hes6/04ZGjAeC7WiYCVXK9cMzK60J18gve/GZNqgqtVxToPf4sEZ/+orgdZ6+9sGqt+AStDb3e8w==;EndpointSuffix=core.windows.net"
container_name = "docs"

# Create the BlobServiceClient object
blob_service_client = BlobServiceClient.from_connection_string(azure_connection_string)
container_client = blob_service_client.get_container_client(container_name)

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

class Chatbot():
    
    def extract_text(self, pdf):
        print("Parsing paper")
        number_of_pages = len(pdf.pages)
        print(f"Total number of pages: {number_of_pages}")
        paper_text = []
        for i in range(number_of_pages):
            page = pdf.pages[i]
            page_text = []

            def visitor_body(text, cm, tm, fontDict, fontSize):
                x = tm[4]
                y = tm[5]
                # ignore header/footer
                if (y > 50 and y < 720) and (len(text.strip()) > 1):
                    page_text.append({
                    'fontsize': fontSize,
                    'text': text.strip().replace('\x03', ''),
                    'x': x,
                    'y': y
                    })

            _ = page.extract_text(visitor_text=visitor_body)

            blob_font_size = None
            blob_text = ''
            processed_text = []

            for t in page_text:
                if t['fontsize'] == blob_font_size:
                    blob_text += f" {t['text']}"
                    if len(blob_text) >= 2000:
                        processed_text.append({
                            'fontsize': blob_font_size,
                            'text': blob_text,
                            'page': i
                        })
                        blob_font_size = None
                        blob_text = ''
                else:
                    if blob_font_size is not None and len(blob_text) >= 1:
                        processed_text.append({
                            'fontsize': blob_font_size,
                            'text': blob_text,
                            'page': i
                        })
                    blob_font_size = t['fontsize']
                    blob_text = t['text']
                paper_text += processed_text
        print("Done parsing paper")
        # print(paper_text)
        return paper_text
    
    # A function that parses through a string and returns a dataframe with the text and the page number. 
    # Split the text into chunks of 2000 characters
    def extract_df(self, text):
        df = pd.DataFrame(columns=['text', 'page'])
        page = 0
        for i in range(0, len(text), 2000):
            df = df.append({'text': text[i:i+2000], 'page': page}, ignore_index=True)
            page += 1
        return df

    # A function that parses through a pdf and returns a dataframe with the text and the page number
    def make_df(self, pdf):
        df = pd.DataFrame(columns=['text', 'page'])
        for i in range(len(pdf.pages)):
            page = pdf.pages[i]
            text = page.extract_text()
            text = text.replace('\n', ' ')
            text = text.strip().replace('\x03', '')
            df = df.append({'text': text, 'page': i}, ignore_index=True)
        return df

    # A function that parses through a docx and returns a dataframe with the text and the paragraph number
    def make_doc(self, doc):
        df = pd.DataFrame(columns=['text', 'paragraph'])
        for i in range(len(doc.paragraphs)):
            paragraph = doc.paragraphs[i]
            text = paragraph.text
            text = text.replace('\n', ' ')
            text = text.strip().replace('\x03', '')
            df = df.append({'text': text, 'paragraph': i}, ignore_index=True)
            # drop row if text is empty
            df = df.drop(df[df['text'] == ''].index)
        return df

    def create_df(self, pdf):
        print('Creating dataframe')
        filtered_pdf= []
        for row in pdf:
            if len(row['text']) < 30:
                continue
            filtered_pdf.append(row)
        df = pd.DataFrame(filtered_pdf)
        # print(df.shape)
        # remove elements with identical df[text] and df[page] values
        df = df.drop_duplicates(subset=['text', 'page'], keep='first')
        df['length'] = df['text'].apply(lambda x: len(x))
        print('Done creating dataframe')
        return df

@app.get('/viewer')
async def viewer(request: Request):
    return templates.TemplateResponse('viewer.html', {"request": request})

@app.get('/app.html')
async def grid(request: Request):
    return templates.TemplateResponse('test.html', {"request": request})

@app.get("/")
@app.post("/")
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

def get_repo_info(url):
    # Extract the repo name, owner, and branch from the URL
    url_parts = url.rstrip("/").split("/")
    current_file = None

    if "blob" in url_parts:
        repo_index = url_parts.index("blob") - 1
        owner, repo_name, branch = url_parts[repo_index - 1], url_parts[repo_index], url_parts[repo_index + 2]
        current_file = "/".join(url_parts[repo_index + 3:])
    else:
        owner, repo_name = url_parts[-2:]
        branch = "main"

    # Define the API URL for the repo
    api_url = f"https://api.github.com/repos/{owner}/{repo_name}/git/trees/{branch}?recursive=1"

    # Send a request to the GitHub API
    response = requests.get(api_url)

    # Raise an exception if the request fails
    response.raise_for_status()

    # Extract the JSON data from the response
    data = response.json()

    # Extract the file paths from the JSON data
    files = [item["path"] for item in data["tree"] if item["type"] == "blob"]

    # Construct the repo URL
    repo_url = f"https://github.com/{owner}/{repo_name}"

    # Construct the file URLs as a dictionary of {file: url}
    file_urls = {file: f"{repo_url}/blob/{branch}/{file}" for file in files}

    # Construct the raw.githubusercontent.com URLs as a dictionary of {file: raw_url}
    raw_urls = {file: f"https://raw.githubusercontent.com/{owner}/{repo_name}/{branch}/{file}" for file in files}

    return {
        "repo_url": repo_url,
        "branch": branch,
        "current_file": current_file,
        "files": files,
        "file_urls": file_urls,
        "raw_urls": raw_urls,
    }

@app.post("/git")
async def git_endpoint(request: Request):
    # Extract the URL from the request
    data = await request.json()
    url = data['url']
    repo_info = get_repo_info(url)
    return repo_info

@app.post('/save_file')
async def save_file(request: Request, model: str = Header(None), content_type: str = Header(None)):
    print('Saving file')
    file = await request.body()
    print(model)
    key = md5(file).hexdigest()
    print(key)

    print(content_type)

    if content_type == 'text/plain' or content_type == 'application/x-javascript' or content_type == 'application/json':
        print('File is text')
        text = file.decode('utf-8')
        # print(text)

        # Create a new Document
        doc = docx.Document()
        paragraphs = text.split('\n')

        # add each paragraph to the document as a new paragraph object
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)

        # Save the Document to an in-memory buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # key = md5(buffer.getvalue()).hexdigest()

        blob_client = container_client.get_blob_client(key)
        content_settings = ContentSettings(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        blob_client.upload_blob(buffer, blob_type="BlockBlob", content_settings=content_settings, overwrite=True)

        print(key)
        return JSONResponse(content={"key": key, "exists": False, "file_type": 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})


    # if blob_exists:
    #     print("File already exists")
    #     return JSONResponse(content={"key": key, "exists": True})

    # If the file doesn't exist, create a new blob client and save the file to Azure Blob Storage
    blob_client = container_client.get_blob_client(key)
    content_settings = ContentSettings(content_type=content_type)
    blob_client.upload_blob(file, blob_type="BlockBlob", content_settings=content_settings, overwrite=True)

    # directory = './files'
    
    # # Check if file exists in directory
    # if os.path.isfile(f'{directory}/{key}'):
    #     print('File already exists')
    #     return JSONResponse(content={"key": key, "exists": True})
    
    # # Save the file to the directory
    # with open(f'{directory}/{key}', 'wb') as f:
    #     f.write(file)

    print('dog')

    print('File uploaded to Azure Blob Storage')

    return JSONResponse(content={"key": key, "exists": False, "file_type": content_type})

@app.get('/get_file/{file_id}')
async def get_file(file_id: str, file_type: str = Header(None)):
    print('Getting file from Azure Blob Storage')
    print(file_id)

    # Get the BlobClient for the file
    container_client = blob_service_client.get_container_client(container_name)
    blob_client = container_client.get_blob_client(file_id)

    if not blob_client.exists():
        print('File does not exist')
        return JSONResponse(content={"error": "File does not exist"})

    # directory = './files'
    
    # # Check if file exists in directory
    # if os.path.isfile(f'{directory}/{file_id}'):
    #     print('File already exists')
    #     return JSONResponse(content={"key": file_id, "exists": True})
    
    # # Get the file from the directory
    # with open(f'{directory}/{file_id}', 'rb') as f:
    #     file = f.read()
    
    # if file_type == 'text/plain':
    #     print('File is pdf')
    #     file_type = 'text/plain'
    #     text = file.decode('utf-8')
    #     chatbot = Chatbot()
    #     df = chatbot.extract_df(text)
    #     print(df)
    #     json_str = df.to_json(orient='records')
    #     return JSONResponse(content={"url": file_id, "file_type": file_type, "df": json_str})

    url = blob_client.url
    print('File url: ', url)

    # Get the file from Azure Blob Storage
    if file_type == 'text/plain':
        print('File is pdf')
        file_type = 'text/plain'
        text = blob_client.download_blob().content_as_text()
        chatbot = Chatbot()
        df = chatbot.extract_df(text)
        # print(df)
        json_str = df.to_json(orient='records')
        return JSONResponse(content={"url": url, "file_type": file_type, "df": json_str})

    # url = ""

    return JSONResponse(content={"url": url, "file_type": file_type})

@app.post("/save")
async def save(request: Request):
    print("Saving df to Azure Blob Storage")

    request_data = await request.json()
    df = request_data['df']
    key = request_data['key']

    df = pd.DataFrame.from_dict(df)

    # Return error if the dataframe is empty
    if df.empty:
        return JSONResponse(content={"error": "No data found"})
    
    # directory = './files'
    
    # # Check if file exists in directory
    # if os.path.isfile(f'{directory}/{key}'):
    #     print('File already exists')
    #     return JSONResponse(content={"key": key, "exists": True})
    
    # # Save the df to the directory
    # with open(f'{directory}/{key}', 'w') as f:
    #     f.write(df.to_json())

    # Get the container client
    container_client = blob_service_client.get_container_client(container_name)
    blob_name = key + '.json'

    # # Check if the file already exists
    # blob_exists = any(blob.name == blob_name for blob in container_client.list_blobs())
    # print(blob_exists)

    # if blob_exists:
    #     print("File already exists")
    #     print("Done processing pdf")
    #     return JSONResponse(content={"key": key, "exists": True})

    # If the file doesn't exist, create a new blob client and upload the file's content
    blob_client = container_client.get_blob_client(blob_name)
    blob_client.upload_blob(df.to_json(), content_type='application/json', overwrite=True)

    print("Saved pdf")
    return JSONResponse(content={"key": key, "exists": False})

@app.post("/get_df")
async def get_df(request: Request):
    print('Getting dataframe')

    request_data = await request.json()
    key = request_data['key']
    print(key)

    if key is None or key == '' or key == 'null' or key == 'undefined':
        print("No key found")
        return JSONResponse(content={"error": "No key found"})

    query = request_data['query']
    print(query)

    # directory = './files'

    # # Check if JSON exists in directory
    # with open(f'{directory}/{key}', 'r') as f:
    #     df = pd.read_json(f.read())

    # Get the container client
    container_client = blob_service_client.get_container_client(container_name)
    blob_name = key + '.json'
    blob_client = container_client.get_blob_client(blob_name)

    # Check if the file exists
    if not blob_client.exists():
        return JSONResponse(content={"error": "File does not exist"})
    
    # Download and read the JSON file
    with BytesIO() as json_buffer:
        blob_client.download_blob().readinto(json_buffer)
        json_buffer.seek(0)
        df = pd.read_json(json_buffer)

    print('Done getting dataframe')
    json_str = df.to_json(orient='records')

    return JSONResponse(content={"df": json_str})

@app.post("/download_pdf")
async def download_pdf(request: Request):
    print("Downloading pdf")

    request_data = await request.json()
    # print(request_data)
    chatbot = Chatbot()
    t = request_data['type']

    key = request_data['key']

    print(key)

    model = request_data['model']

    if t == 'application/pdf':
        print('File is pdf')
        file_type = 'application/pdf'

        data = request_data['data']

        # Decode the base64 string
        data = base64.b64decode(data)

        key = md5(data).hexdigest()
        print(key)
        pdf = PdfReader(BytesIO(data))

        if model == 'gpt-4':
            df = chatbot.make_df(pdf)
        else:
            paper_text = chatbot.extract_text(pdf)
            df = chatbot.create_df(paper_text)
        json_str = df.to_json(orient='records')

    elif t == 'application/msword' or t == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        url = request_data['url']
        print(url)
        r = requests.get(url, allow_redirects=True, headers={
        "Origin": "appspot.com", 'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, Authorization'})

        file = r.content

        doc = docx.Document(BytesIO(file))

        chatbot = Chatbot()
        df = chatbot.make_doc(doc)
        json_str = df.to_json(orient='records')

    else:
        url = request_data['url']
        print(url)
        r = requests.get(url, allow_redirects=True, headers={
        "Origin": "appspot.com", 'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, Authorization'})

        file = r.content

        def read_txt_file(file):
            paragraphs = file.split('\n').split('\r').replace('\n', '').replace('\r', '').replace('\t', '').replace('&#160;', '')
            df = pd.DataFrame(paragraphs, columns=['text'])
            # print(df)
            df['paragraph_number'] = df.index
            return df

        df = read_txt_file(file.decode('utf-8'))

        json_str = df.to_json(orient='records')

    # Get the container client
    container_client = blob_service_client.get_container_client(container_name)
    blob_name = key + '.json'

    blob_exists = False

    print(blob_name)
    print('duck')

    for blob in container_client.list_blobs():
        if str(blob.name+'.json') == str(blob_name+'.json'):
            print('blob exists')
            blob_exists = True
            break

    print('moose')

    print(blob_exists)

    if blob_exists:
        print("File already exists")
        print("Done processing pdf")
        return JSONResponse(content={"key": key, "exists": True})

    # If the file doesn't exist, create a new blob client and upload the JSON file to Azure Blob Storage
    blob_client = container_client.get_blob_client(blob_name)
    print('saved to blob storage')
    blob_client.upload_blob(json_str, overwrite=True)


    # directory = './files'

    # # Save the df to the directory
    # with open(f'{directory}/{key}', 'w') as f:
    #     f.write(json_str)

    print("Done processing pdf")
    return JSONResponse(content={"key": key, "exists": False, "df": json_str})
